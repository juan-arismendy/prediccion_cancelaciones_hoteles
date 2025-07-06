#!/usr/bin/env python3
"""
Script para extraer el contenido del archivo Reporte.docx
y convertirlo a texto plano para análisis
"""

import docx
import sys

def extract_docx_content(docx_path, output_path=None):
    """Extrae el contenido de un archivo DOCX"""
    try:
        # Abrir el documento
        doc = docx.Document(docx_path)
        
        # Extraer todo el texto
        full_text = []
        
        print("📄 EXTRAYENDO CONTENIDO DEL REPORTE...")
        print("=" * 50)
        
        # Extraer párrafos
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():  # Solo párrafos con contenido
                full_text.append(paragraph.text)
        
        # Extraer tablas
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(" | ".join(row_text))
        
        # Unir todo el texto
        complete_text = "\n".join(full_text)
        
        # Guardar en archivo si se especifica
        if output_path:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(complete_text)
            print(f"✅ Contenido guardado en: {output_path}")
        
        # Mostrar estadísticas
        print(f"\n📊 ESTADÍSTICAS DEL DOCUMENTO:")
        print(f"   - Total párrafos: {len(doc.paragraphs)}")
        print(f"   - Total tablas: {len(doc.tables)}")
        print(f"   - Total caracteres: {len(complete_text)}")
        print(f"   - Total palabras: {len(complete_text.split())}")
        
        # Mostrar preview
        print(f"\n📖 PREVIEW DEL CONTENIDO:")
        print("-" * 50)
        preview = complete_text[:2000] + "..." if len(complete_text) > 2000 else complete_text
        print(preview)
        
        return complete_text
        
    except Exception as e:
        print(f"❌ Error al procesar el archivo: {e}")
        return None

if __name__ == "__main__":
    docx_file = "datos/Reporte.docx"
    output_file = "reporte_extracted.txt"
    
    print("🔍 EXTRACTOR DE CONTENIDO DOCX")
    print("=" * 50)
    
    content = extract_docx_content(docx_file, output_file)
    
    if content:
        print(f"\n✅ EXTRACCIÓN COMPLETADA")
        print(f"📁 Archivo original: {docx_file}")
        print(f"📄 Archivo de texto: {output_file}")
        print(f"\n💡 Ahora puedo analizar el contenido del reporte!")
    else:
        print(f"\n❌ No se pudo extraer el contenido") 